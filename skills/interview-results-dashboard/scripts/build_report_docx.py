from __future__ import annotations

import html
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "07_report_draft.md"
OUTPUT_DOCX = ROOT / "08_report_draft.docx"
OUTPUT_HTML = ROOT / "09_report_preview.html"

TITLE = "AI연금투자솔루션 사용자 리서치 결과 보고서"
VERSION_LABEL = "v0.1_260818"
META_LINE = "2026.08  |  하나금융융합기술원 HXR Cell"

BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Malgun Gothic"
BODY_COLOR = RGBColor(33, 40, 52)
MUTED = RGBColor(98, 108, 125)
H1_COLOR = RGBColor(46, 116, 181)
H3_COLOR = RGBColor(31, 77, 120)
ACCENT = RGBColor(14, 101, 124)


def set_run_font(
    run,
    *,
    name: str = BODY_FONT,
    east_asia: str = EAST_ASIA_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color: str = "D9E3E8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    h1.font.name = BODY_FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    h1.font.size = Pt(16)
    h1.font.color.rgb = H1_COLOR
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.10

    h2 = doc.styles["Heading 2"]
    h2.font.name = BODY_FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    h2.font.size = Pt(13)
    h2.font.color.rgb = H1_COLOR
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.10

    h3 = doc.styles["Heading 3"]
    h3.font.name = BODY_FONT
    h3._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    h3.font.size = Pt(12)
    h3.font.color.rgb = H3_COLOR
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(11)
        style.font.color.rgb = BODY_COLOR
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.10


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_footer(section) -> None:
    section.different_first_page_header_footer = True

    first_footer = section.first_page_footer
    if first_footer.paragraphs:
        first_footer.paragraphs[0].clear()

    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(TITLE)
    set_run_font(run, size=9, color=MUTED)


def add_header(section) -> None:
    first_header = section.first_page_header
    if first_header.paragraphs:
        first_header.paragraphs[0].clear()

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("HX Research Cell")
    set_run_font(run, size=9, color=MUTED)


def add_cover_page(doc: Document) -> None:
    version_p = doc.add_paragraph()
    version_p.paragraph_format.space_before = Pt(140)
    version_p.paragraph_format.space_after = Pt(4)
    run = version_p.add_run(VERSION_LABEL)
    set_run_font(run, size=10.5, color=MUTED)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(10)
    run = title_p.add_run(TITLE)
    set_run_font(run, size=24, color=RGBColor(25, 56, 82), bold=True)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(24)
    run = meta_p.add_run(META_LINE)
    set_run_font(run, size=11, color=ACCENT)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(0)
    set_paragraph_border_bottom(rule)

    doc.add_page_break()


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(f"▍ {text}")
    set_run_font(run, size=10.5, color=RGBColor(74, 83, 98), italic=True)


def add_quote_attribution(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"— {text}")
    set_run_font(run, size=10, color=MUTED)


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        pass

    header_cells = table.rows[0].cells
    for idx, value in enumerate(header):
        cell = header_cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(value)
        set_run_font(run, size=10.5, color=RGBColor(255, 255, 255), bold=True)
        cell._tc.get_or_add_tcPr()
        shade = OxmlElement("w:shd")
        shade.set(qn("w:val"), "clear")
        shade.set(qn("w:fill"), "2E74B5")
        cell._tc.tcPr.append(shade)

    for row in body:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, size=10.5, color=BODY_COLOR)

    doc.add_paragraph()


def parse_table_block(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in lines:
        parts = [part.strip() for part in raw.strip().strip("|").split("|")]
        if parts and all(set(part) <= {"-", ":"} for part in parts):
            continue
        rows.append(parts)
    return rows


def parse_markdown(md_text: str) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.rstrip()

        if not line:
            items.append(("blank", ""))
            i += 1
            continue

        if line.lstrip().startswith("<!--"):
            i += 1
            continue

        if line.lstrip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].rstrip().lstrip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            items.append(("table", parse_table_block(table_lines)))
            continue

        if line.startswith("▍ "):
            items.append(("quote", line[2:].strip()))
            i += 1
            continue

        if line.startswith("— "):
            items.append(("quote_attr", line[2:].strip()))
            i += 1
            continue

        if line.startswith("#### "):
            items.append(("h4", line[5:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            items.append(("h3", line[4:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            items.append(("h2", line[3:].strip()))
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("- "):
            items.append(("bullet", line[2:].strip()))
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            items.append(("bullet", re.sub(r"^\d+\.\s*", "", line).strip()))
            i += 1
            continue

        items.append(("p", line.strip()))
        i += 1

    return items


def build_body(doc: Document, items: list[tuple[str, str]]) -> None:
    first_h1 = True
    for kind, payload in items:
        if kind == "blank":
            continue
        if kind == "h2":
            if not first_h1:
                doc.add_page_break()
            p = doc.add_paragraph(payload, style="Heading 1")
            if first_h1:
                p.paragraph_format.space_before = Pt(0)
                first_h1 = False
            continue
        if kind == "h3":
            doc.add_paragraph(payload, style="Heading 2")
            continue
        if kind == "h4":
            doc.add_paragraph(payload, style="Heading 3")
            continue
        if kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(payload)
            set_run_font(run, size=11, color=BODY_COLOR)
            continue
        if kind == "quote":
            add_quote(doc, payload)
            continue
        if kind == "quote_attr":
            add_quote_attribution(doc, payload)
            continue
        if kind == "table":
            add_docx_table(doc, payload)
            continue
        p = doc.add_paragraph()
        run = p.add_run(payload)
        set_run_font(run, size=11, color=BODY_COLOR)


def render_quote_html(text: str) -> str:
    return f'<p class="quote-text">▍ {html.escape(text)}</p>'


def render_table_html(rows: list[list[str]]) -> str:
    if len(rows) < 2:
        return ""
    header, body = rows[0], rows[1:]
    head_html = "".join(f"<th>{html.escape(cell)}</th>" for cell in header)
    body_html = "".join(
        "<tr>" + "".join(f"<td>{html.escape(cell)}</td>" for cell in row) + "</tr>" for row in body
    )
    return f'<table class="report-table"><thead><tr>{head_html}</tr></thead><tbody>{body_html}</tbody></table>'


def build_html_preview(items: list[tuple[str, str]]) -> Path:
    body_parts: list[str] = []
    current_section_open = False
    list_open = False
    quote_open = False

    def close_list() -> None:
        nonlocal list_open
        if list_open:
            body_parts.append("</ul>")
            list_open = False

    def close_quote() -> None:
        nonlocal quote_open
        if quote_open:
            body_parts.append("</blockquote>")
            quote_open = False

    def close_section() -> None:
        nonlocal current_section_open
        close_list()
        close_quote()
        if current_section_open:
            body_parts.append("</section>")
            current_section_open = False

    for kind, payload in items:
        if kind == "blank":
            close_list()
            continue
        if kind == "h2":
            close_section()
            body_parts.append('<section class="report-section">')
            body_parts.append(f"<h2>{html.escape(payload)}</h2>")
            current_section_open = True
            continue
        if kind == "h3":
            close_list()
            close_quote()
            body_parts.append(f"<h3>{html.escape(payload)}</h3>")
            continue
        if kind == "h4":
            close_list()
            close_quote()
            body_parts.append(f'<h4 class="insight-title">{html.escape(payload)}</h4>')
            continue
        if kind == "bullet":
            if not list_open:
                body_parts.append('<ul class="bullet-list">')
                list_open = True
            body_parts.append(f"<li>{html.escape(payload)}</li>")
            continue
        close_list()
        if kind == "table":
            close_quote()
            body_parts.append(render_table_html(payload))
            continue
        if kind == "quote":
            close_quote()
            body_parts.append('<blockquote class="quote">')
            quote_open = True
            body_parts.append(render_quote_html(payload))
            continue
        if kind == "quote_attr":
            body_parts.append(f'<div class="quote-attr">— {html.escape(payload)}</div>')
            close_quote()
            continue
        body_parts.append(f"<p>{html.escape(payload)}</p>")

    close_section()

    html_text = f"""<!DOCTYPE html>
<html lang="ko">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(TITLE)} - Preview</title>
  <style>
    :root {{
      --bg: #e8eef1;
      --paper: #ffffff;
      --ink: #21303b;
      --muted: #5f6e7d;
      --line: #d9e3e8;
      --accent: #2e74b5;
      --accent-strong: #193852;
      --quote-bg: #f4f8fb;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      background:
        radial-gradient(circle at top left, rgba(46, 116, 181, 0.12), transparent 30%),
        linear-gradient(180deg, #f3f7f8 0%, var(--bg) 100%);
      color: var(--ink);
      font-family: "Segoe UI", "Malgun Gothic", sans-serif;
      line-height: 1.72;
    }}
    .shell {{
      max-width: 1040px;
      margin: 0 auto;
      padding: 40px 20px 80px;
    }}
    .cover {{
      background: linear-gradient(180deg, #ffffff 0%, #f6fafc 100%);
      border: 1px solid rgba(46, 116, 181, 0.16);
      border-radius: 28px;
      padding: 56px 56px 48px;
      box-shadow: 0 24px 64px rgba(26, 53, 71, 0.08);
    }}
    .eyebrow {{
      display: inline-block;
      margin-bottom: 18px;
      padding: 6px 12px;
      border-radius: 999px;
      background: rgba(46, 116, 181, 0.1);
      color: var(--accent);
      font-size: 13px;
      font-weight: 700;
      letter-spacing: 0.04em;
    }}
    h1 {{
      margin: 0 0 10px;
      color: var(--accent-strong);
      font-size: clamp(32px, 4vw, 44px);
      line-height: 1.24;
      letter-spacing: -0.02em;
    }}
    .meta {{
      margin-top: 20px;
      padding-top: 20px;
      border-top: 1px solid var(--line);
      color: var(--muted);
      font-size: 14px;
    }}
    .report-section {{
      margin-top: 28px;
      padding: 40px 48px;
      background: var(--paper);
      border: 1px solid rgba(46, 116, 181, 0.14);
      border-radius: 24px;
      box-shadow: 0 16px 42px rgba(30, 50, 66, 0.06);
    }}
    h2 {{
      margin: 0 0 18px;
      color: var(--accent);
      font-size: 29px;
      letter-spacing: -0.02em;
    }}
    h3 {{
      margin: 34px 0 10px;
      color: var(--accent);
      font-size: 21px;
    }}
    .insight-title {{
      margin: 24px 0 10px;
      color: #1f4d78;
      font-size: 18px;
      line-height: 1.5;
    }}
    p {{
      margin: 0 0 14px;
      font-size: 16px;
    }}
    .bullet-list {{
      margin: 0 0 16px;
      padding-left: 22px;
    }}
    .bullet-list li {{
      margin-bottom: 8px;
    }}
    .report-table {{
      width: 100%;
      border-collapse: collapse;
      margin: 0 0 20px;
      font-size: 14px;
    }}
    .report-table th, .report-table td {{
      border: 1px solid var(--line);
      padding: 8px 10px;
      text-align: left;
      vertical-align: top;
    }}
    .report-table th {{
      background: rgba(46, 116, 181, 0.12);
      color: var(--accent-strong);
    }}
    .quote {{
      margin: 18px 0 22px;
      padding: 18px 22px;
      border-radius: 18px;
      border: 1px solid rgba(46, 116, 181, 0.12);
      background: var(--quote-bg);
    }}
    .quote-text {{
      margin: 0 0 6px;
      color: #4a5362;
      font-style: italic;
      font-size: 15px;
    }}
    .quote-attr {{
      margin: -12px 0 22px;
      padding: 0 22px;
      color: var(--muted);
      font-size: 13px;
    }}
    @media (max-width: 760px) {{
      .shell {{ padding: 20px 14px 56px; }}
      .cover, .report-section {{ padding: 28px 22px; border-radius: 20px; }}
      p, .bullet-list li {{ font-size: 15px; }}
    }}
  </style>
</head>
<body>
  <main class="shell">
    <section class="cover">
      <div class="eyebrow">{html.escape(VERSION_LABEL)}</div>
      <h1>{html.escape(TITLE)}</h1>
      <div class="meta">{html.escape(META_LINE)}</div>
    </section>
    {"".join(body_parts)}
  </main>
</body>
</html>
"""

    OUTPUT_HTML.write_text(html_text, encoding="utf-8")
    return OUTPUT_HTML


def build_docx() -> Path:
    md_text = SOURCE_MD.read_text(encoding="utf-8")
    items = parse_markdown(md_text)
    doc = Document()

    configure_section(doc.sections[0])
    configure_styles(doc)
    add_header(doc.sections[0])
    add_footer(doc.sections[0])

    add_cover_page(doc)
    build_body(doc, items)

    doc.save(OUTPUT_DOCX)
    build_html_preview(items)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = build_docx()
    print(path)
    print(OUTPUT_HTML)
