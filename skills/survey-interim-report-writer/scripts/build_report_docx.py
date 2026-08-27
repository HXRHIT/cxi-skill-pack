from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


VERSION_AND_TITLE_PATTERN = re.compile(r"^(v[\d.]+_\d{6})\s+(.+)$")
VERSION_LABEL_PATTERN = re.compile(r"^v[\d.]+_\d{6}(\s.*)?$")
BULLET_PREFIX = "\u2022 "
DEFAULT_METADATA = "2026.08  |  \ud558\ub098\uae08\uc735\uc735\ud569\uae30\uc220\uc6d0 HXR Cell"
LATEST_NOTE = "\ucd5c\uc2e0\ud310 \ucd94\uac00 \ube14\ub85d | 2026-08-18 | survey-interim-report-writer validation"


def parse_markdown(path: Path) -> list[tuple]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[tuple] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        if line.lstrip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            blocks.append(("table", table_lines))
            continue

        if line.startswith("- "):
            items = []
            while i < len(lines) and lines[i].startswith("- "):
                items.append(lines[i][2:].strip())
                i += 1
            blocks.append(("bullets", items))
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            blocks.append(("heading", len(heading.group(1)), heading.group(2).strip()))
            i += 1
            continue

        paragraph_lines = [line.strip()]
        i += 1
        while i < len(lines):
            candidate = lines[i].rstrip()
            if not candidate.strip():
                i += 1
                break
            if candidate.lstrip().startswith("|") or candidate.startswith("- ") or re.match(
                r"^(#{1,3})\s+", candidate
            ):
                break
            paragraph_lines.append(candidate.strip())
            i += 1
        blocks.append(("paragraph", " ".join(paragraph_lines)))

    return blocks


def clean_inline_text(text: str) -> str:
    cleaned = text.replace("`", "").replace("**", "")
    return cleaned.strip()


def extract_title_block(blocks: list[tuple]) -> tuple[str, str, str | None, str | None, list[tuple]]:
    version_label = "v1.0_260818"
    report_title = "Survey Interim Report Draft"
    metadata_line: str | None = None
    note_text: str | None = None
    remaining = list(blocks)

    # House format: a bare version-label paragraph, then a level-1 heading title.
    if (
        remaining
        and remaining[0][0] == "paragraph"
        and VERSION_LABEL_PATTERN.match(clean_inline_text(remaining[0][1]))
    ):
        version_label = clean_inline_text(remaining[0][1])
        remaining = remaining[1:]

    if remaining and remaining[0][0] == "heading" and remaining[0][1] == 1:
        heading_text = clean_inline_text(remaining[0][2])
        match = VERSION_AND_TITLE_PATTERN.match(heading_text)
        if match:
            # Backward-compatible combined "vX.Y_yymmdd 제목" heading form.
            version_label, report_title = match.groups()
        else:
            report_title = heading_text
        remaining = remaining[1:]

    # House format: a bare "yyyy.mm | 조직명" metadata line right under the title.
    if (
        remaining
        and remaining[0][0] == "paragraph"
        and "|" in clean_inline_text(remaining[0][1])
        and "validation" not in remaining[0][1].lower()
    ):
        metadata_line = clean_inline_text(remaining[0][1])
        remaining = remaining[1:]

    if remaining and remaining[0][0] == "paragraph":
        first_paragraph = clean_inline_text(remaining[0][1])
        if (
            "validation" in first_paragraph.lower()
            or "survey-interim-report-writer" in first_paragraph.lower()
            or "dashboard-handoff" in first_paragraph.lower()
        ):
            note_text = first_paragraph
            remaining = remaining[1:]

    return version_label, report_title, metadata_line, note_text, remaining


def get_style_name(doc: Document, *candidates: str) -> str | None:
    for candidate in candidates:
        try:
            doc.styles[candidate]
            return candidate
        except KeyError:
            continue
    return None


def extract_metadata_line(doc: Document) -> str:
    nonempty = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    if len(nonempty) >= 3:
        return nonempty[2]
    return DEFAULT_METADATA


def extract_table_style_id(doc: Document) -> str | None:
    if not doc.tables:
        return None
    tbl_style = doc.tables[0]._tbl.tblPr.find(qn("w:tblStyle"))
    if tbl_style is None:
        return None
    return tbl_style.get(qn("w:val"))


def normalize_section_margin_xml(doc: Document) -> None:
    for section in doc.sections:
        sect_pr = section._sectPr
        pg_mar = sect_pr.pgMar
        if pg_mar is None:
            continue
        for attr in ("left", "right", "top", "bottom", "header", "footer", "gutter"):
            value = pg_mar.get(qn(f"w:{attr}"))
            if not value:
                continue
            try:
                int(value)
            except ValueError:
                try:
                    pg_mar.set(qn(f"w:{attr}"), str(int(round(float(value)))))
                except ValueError:
                    continue


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)


def load_template_copy(base_docx: Path, out_path: Path) -> tuple[Document, Path]:
    working_copy = out_path.with_name(f"{out_path.stem}__working_copy.docx")
    shutil.copy2(base_docx, working_copy)
    doc = Document(working_copy)
    normalize_section_margin_xml(doc)
    return doc, working_copy


def add_styled_paragraph(
    doc: Document, text: str = "", style_candidates: tuple[str, ...] = ("normal", "Normal")
):
    style_name = get_style_name(doc, *style_candidates)
    if style_name:
        paragraph = doc.add_paragraph(style=style_name)
    else:
        paragraph = doc.add_paragraph()
    if text:
        paragraph.add_run(text)
    return paragraph


def add_house_title_block(doc: Document, version_label: str, report_title: str, note_text: str | None, metadata: str) -> None:
    add_styled_paragraph(doc, version_label, ("Title", "Heading 1", "normal", "Normal"))
    add_styled_paragraph(doc, report_title, ("Heading 1", "Heading 2", "normal", "Normal"))
    add_styled_paragraph(doc, metadata, ("normal", "Normal"))

    if note_text:
        note_paragraph = add_styled_paragraph(doc, "", ("normal", "Normal"))
        note_paragraph.add_run(note_text).italic = True


def add_heading(doc: Document, level: int, text: str) -> None:
    style_name = get_style_name(doc, f"Heading {min(level, 3)}", "Heading 1", "normal", "Normal")
    if style_name:
        doc.add_paragraph(clean_inline_text(text), style=style_name)
    else:
        doc.add_paragraph(clean_inline_text(text))


def add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
    paragraph = add_styled_paragraph(doc, "", ("normal", "Normal"))
    run = paragraph.add_run(clean_inline_text(text))
    if italic:
        run.italic = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(doc, f"{BULLET_PREFIX}{clean_inline_text(item)}")


def add_table(doc: Document, lines: list[str], table_style_id: str | None) -> None:
    rows = []
    for raw in lines:
        parts = [clean_inline_text(part.strip()) for part in raw.strip().strip("|").split("|")]
        if parts and all(set(part) <= {"-", ":"} for part in parts):
            continue
        rows.append(parts)

    if len(rows) < 2:
        return

    header = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    if table_style_id:
        try:
            table.style = table_style_id
        except (KeyError, ValueError):
            pass

    header_cells = table.rows[0].cells
    for idx, value in enumerate(header):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(value)
        run.bold = True

    for row in body:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    doc.add_paragraph()


def render_blocks(doc: Document, blocks: list[tuple], table_style_id: str | None) -> None:
    for block in blocks:
        kind = block[0]
        if kind == "heading":
            _, level, text = block
            add_heading(doc, level, text)
        elif kind == "paragraph":
            add_paragraph(doc, block[1])
        elif kind == "bullets":
            add_bullets(doc, block[1])
        elif kind == "table":
            add_table(doc, block[1], table_style_id)


def build_standalone(markdown_path: Path, base_docx: Path, out_path: Path) -> None:
    doc, working_copy = load_template_copy(base_docx, out_path)
    template_metadata = extract_metadata_line(doc)
    table_style_id = extract_table_style_id(doc)
    version_label, report_title, metadata_line, note_text, blocks = extract_title_block(
        parse_markdown(markdown_path)
    )
    metadata = metadata_line or template_metadata

    clear_document_body(doc)
    add_house_title_block(doc, version_label, report_title, note_text, metadata)
    render_blocks(doc, blocks, table_style_id)
    doc.save(out_path)
    working_copy.unlink(missing_ok=True)


def build_appended(markdown_path: Path, base_docx: Path, out_path: Path) -> None:
    doc, working_copy = load_template_copy(base_docx, out_path)
    template_metadata = extract_metadata_line(doc)
    table_style_id = extract_table_style_id(doc)
    version_label, report_title, metadata_line, note_text, blocks = extract_title_block(
        parse_markdown(markdown_path)
    )
    metadata = metadata_line or template_metadata

    doc.add_page_break()
    add_paragraph(doc, LATEST_NOTE, italic=True)
    add_house_title_block(doc, version_label, report_title, note_text, metadata)
    render_blocks(doc, blocks, table_style_id)
    doc.save(out_path)
    working_copy.unlink(missing_ok=True)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", required=True)
    parser.add_argument("--standalone", required=True)
    parser.add_argument("--base-docx", required=True)
    parser.add_argument("--appended", required=True)
    args = parser.parse_args()

    markdown_path = Path(args.markdown)
    standalone_path = Path(args.standalone)
    base_docx = Path(args.base_docx)
    appended_path = Path(args.appended)

    standalone_path.parent.mkdir(parents=True, exist_ok=True)
    appended_path.parent.mkdir(parents=True, exist_ok=True)

    build_standalone(markdown_path, base_docx, standalone_path)
    build_appended(markdown_path, base_docx, appended_path)


if __name__ == "__main__":
    main()
